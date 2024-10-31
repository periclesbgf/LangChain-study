# app/database/workspace_handler.py
from typing import Optional, List
import uuid
from datetime import datetime, timezone
from bson import ObjectId
from api.endpoints.models import Material, AccessLevel
from agent.image_handler import ImageHandler
from database.mongo_database_manager import MongoDatabaseManager
from database.vector_db import QdrantHandler, TextSplitter
from pydantic import BaseModel
import io
from PIL import Image
import hashlib
from enum import Enum
import io
import fitz
import pdfplumber
import pandas as pd
from PIL import Image
import hashlib
import pytesseract
import docx
from qdrant_client import models

class WorkspaceHandler:
    def __init__(self, mongo_manager: MongoDatabaseManager, qdrant_handler: QdrantHandler, image_handler: ImageHandler, text_splitter:TextSplitter):
        self.mongo_manager = mongo_manager
        self.qdrant_handler = qdrant_handler
        self.image_handler = image_handler
        self.text_splitter = text_splitter
        self.images_collection = self.mongo_manager.db['image_collection']


    async def add_material(
        self,
        file_content: bytes,
        filename: str,
        student_email: str,
        access_level: AccessLevel,
        discipline_id: Optional[str] = None,
        session_id: Optional[str] = None
    ) -> Material:
        """Adiciona um novo material com nível de acesso especificado."""
        
        # Gera IDs únicos e hash
        content_hash = hashlib.sha256(file_content).hexdigest()
        file_type = self._get_file_type(filename)
        image_uuid = str(uuid.uuid4()) if file_type == 'image' else None

        # Se for imagem, armazena no MongoDB
        if file_type == 'image':
            await self._store_image_in_mongodb(image_uuid, file_content)
            
            # Processa imagem para descrição
            img_base64 = self.image_handler.encode_image_bytes(file_content)
            description = self.image_handler.image_summarize(img_base64)
            
            # Metadata para o Qdrant
            metadata = {
                "student_email": student_email,
                "session_id": session_id,
                "content_hash": content_hash,
                "access_level": access_level.lower(),
                "type": file_type,
                "image_uuid": image_uuid,
                "disciplina": discipline_id
            }
            
            # Adiciona ao Qdrant
            self.qdrant_handler.add_document(
                student_email=student_email,
                disciplina=discipline_id,
                content=description,
                metadata_extra=metadata
            )
        else:
            # Processa documento texto
            text_content = await self._extract_text_content(file_content, file_type)
            
            # Metadata para o Qdrant
            metadata = {
                "student_email": student_email,
                "session_id": session_id,
                "content_hash": content_hash,
                "access_level": access_level.lower(),
                "type": file_type,
                "disciplina": discipline_id
            }
            
            # Adiciona ao Qdrant
            self.qdrant_handler.add_document(
                student_email=student_email,
                disciplina=discipline_id,
                content=text_content,
                metadata_extra=metadata
            )

        # Retorna Material para compatibilidade com API
        return Material(
            id=image_uuid or content_hash,  # Usa image_uuid para imagens, content_hash para outros
            name=filename,
            type=file_type,
            access_level=access_level,
            discipline_id=discipline_id,
            session_id=session_id,
            student_email=student_email,
            content_hash=content_hash,
            created_at=datetime.now(timezone.utc),
            updated_at=datetime.now(timezone.utc)
        )
    def _get_file_type(self, filename: str) -> str:
        """Determina o tipo do arquivo pelo nome."""
        lower_filename = filename.lower()
        if lower_filename.endswith('.pdf'):
            return 'pdf'
        elif lower_filename.endswith(('.jpg', '.jpeg', '.png')):
            return 'image'
        elif lower_filename.endswith(('.doc', '.docx')):
            return 'doc'
        else:
            raise ValueError("Tipo de arquivo não suportado")

    async def _store_in_mongodb(self, material: Material, content: bytes):
        """Armazena metadados do material e conteúdo no MongoDB."""
        document = {
            "_id": material.id,
            "name": material.name,
            "type": material.type,
            "access_level": material.access_level,
            "discipline_id": material.discipline_id,
            "session_id": material.session_id,
            "student_email": material.student_email,
            "content_hash": material.content_hash,
            "content": content,
            "created_at": material.created_at,
            "updated_at": material.updated_at
        }
        await self.mongo_manager.db['student_learn_preference'].insert_one(document)

    async def _process_and_store_in_qdrant(self, material: Material, content: bytes):
        """Processa o conteúdo do arquivo e armazena no Qdrant baseado no tipo."""
        try:
            if material.type == 'image':
                # Gera imagem UUID
                image_uuid = str(uuid.uuid4())
                
                # Armazena imagem no MongoDB
                await self._store_image_in_mongodb(
                    image_uuid=image_uuid,
                    content=content,
                    student_email=material.student_email
                )
                
                # Gera descrição da imagem
                img_base64 = self.image_handler.encode_image_bytes(content)
                description = self.image_handler.image_summarize(img_base64)
                
                # Prepara metadata para Qdrant
                metadata = {
                    "student_email": material.student_email,
                    "session_id": material.session_id,
                    "content_hash": material.content_hash,
                    "access_level": material.access_level.lower(),
                    "type": material.type,
                    "image_uuid": image_uuid,
                    "disciplina": material.discipline_id
                }
                
                # Adiciona ao Qdrant
                self.qdrant_handler.add_document(
                    student_email=material.student_email,
                    disciplina=material.discipline_id,
                    content=description,
                    metadata_extra=metadata
                )
            else:
                # Processa outros tipos de documentos (PDF, DOC)
                text_content = await self._extract_text_content(content, material.type)
                
                # Metadata para outros tipos de documentos
                metadata = {
                    "student_email": material.student_email,
                    "session_id": material.session_id,
                    "content_hash": material.content_hash,
                    "access_level": material.access_level.lower(),
                    "type": material.type,
                    "disciplina": material.discipline_id
                }
                
                # Adiciona ao Qdrant
                self.qdrant_handler.add_document(
                    student_email=material.student_email,
                    disciplina=material.discipline_id,
                    content=text_content,
                    metadata_extra=metadata
                )
                
        except Exception as e:
            print(f"[ERROR] Erro ao processar material {material.id}: {str(e)}")
            raise
    
    async def get_materials(
        self,
        student_email: str,
        discipline_id: Optional[str] = None,
        session_id: Optional[str] = None
    ) -> List[Material]:
        """
        Recupera materiais baseado no nível de acesso e contexto.
        Retorna materiais que são:
        1. Globais (acessíveis a todas as disciplinas)
        2. Específicos da disciplina (se discipline_id for fornecido)
        3. Específicos da sessão (se session_id for fornecido)
        """
        query = {"student_email": student_email}
        
        if session_id:
            # Para contexto de sessão, obtém:
            # - Materiais globais
            # - Materiais da disciplina específica
            # - Materiais da sessão específica
            query["$or"] = [
                {"access_level": AccessLevel.GLOBAL},
                {
                    "access_level": AccessLevel.DISCIPLINE,
                    "discipline_id": discipline_id
                },
                {
                    "access_level": AccessLevel.SESSION,
                    "session_id": session_id
                }
            ]
        elif discipline_id:
            # Para contexto de disciplina, obtém:
            # - Materiais globais
            # - Materiais da disciplina específica
            query["$or"] = [
                {"access_level": AccessLevel.GLOBAL},
                {
                    "access_level": AccessLevel.DISCIPLINE,
                    "discipline_id": discipline_id
                }
            ]
        else:
            # Apenas materiais globais
            query["access_level"] = AccessLevel.GLOBAL

        cursor = self.mongo_manager.db['student_learn_preference'].find(query)
        materials = []
        async for doc in cursor:
            materials.append(Material(**doc))
        return materials

    async def _store_image_in_mongodb(self, image_uuid: str, content: bytes, student_email: str):
        """Armazena imagem no MongoDB usando o formato correto."""
        document = {
            "_id": image_uuid,
            "student_email": student_email,
            "image_data": content,
            "timestamp": datetime.now(timezone.utc)
        }
        await self.images_collection.insert_one(document)

    async def get_material_content(self, material_id: str) -> Optional[bytes]:
        """Recupera o conteúdo da imagem do MongoDB."""
        try:
            # Busca no MongoDB usando o UUID
            image_doc = await self.images_collection.find_one({"_id": material_id})
            if image_doc and 'image_data' in image_doc:
                return image_doc['image_data']
            return None
        except Exception as e:
            print(f"[ERROR] Erro ao recuperar imagem {material_id}: {str(e)}")
            return None

    async def update_material_access(
        self,
        material_id: str,
        new_access_level: AccessLevel,
        discipline_id: Optional[str] = None,
        session_id: Optional[str] = None
    ):
        """Atualiza o nível de acesso do material e contexto."""
        # Atualiza no MongoDB
        update_data = {
            "access_level": new_access_level,
            "discipline_id": discipline_id,
            "session_id": session_id,
            "updated_at": datetime.now(timezone.utc)
        }
        
        await self.mongo_manager.db['student_learn_preference'].update_one(
            {"_id": material_id},
            {"$set": update_data}
        )

        # Atualiza no Qdrant
        metadata_update = {
            "access_level": new_access_level,
            "discipline_id": discipline_id,
            "session_id": session_id
        }
        
        # Recupera o documento atual do MongoDB para obter todos os dados necessários
        doc = await self.mongo_manager.db['student_learn_preference'].find_one({"_id": material_id})
        if doc:
            self.qdrant_handler.add_document(
                student_email=doc["student_email"],
                disciplina=discipline_id or doc["discipline_id"],
                content=doc["content"],
                metadata_extra=metadata_update
            )

    async def delete_material(self, material_id: str):
        """Remove o material do MongoDB e Qdrant."""
        try:
            print(f"[DELETE] Iniciando remoção do material: {material_id}")
            
            # Remove do Qdrant primeiro para manter a referência à imagem
            query_filter = models.Filter(
                must=[
                    models.FieldCondition(
                        key="metadata.image_uuid",
                        match=models.MatchValue(value=material_id)
                    )
                ]
            )
            
            # Remove do Qdrant
            self.qdrant_handler.delete(
                collection_name=self.qdrant_handler.collection_name,
                points_filter=query_filter
            )
            
            # Remove a imagem do MongoDB se existir
            result = await self.images_collection.delete_one({"_id": material_id})
            if result.deleted_count > 0:
                print(f"[DELETE] Imagem removida do MongoDB: {material_id}")
            
        except Exception as e:
            print(f"[ERROR] Erro ao remover material {material_id}: {str(e)}")
            raise

    async def _extract_text_content(self, content: bytes, file_type: str, material_id: str) -> str:
        """
        Extrai e processa o conteúdo do arquivo baseado no tipo.
        
        Args:
            content: Bytes do arquivo
            file_type: Tipo do arquivo ('pdf', 'doc', 'image')
            material_id: ID do material para referência
            
        Returns:
            str: Texto extraído e processado
        """
        try:
            extracted_content = ""
            
            if file_type == 'pdf':
                # Processa PDF
                pdf_document = fitz.open(stream=content, filetype="pdf")
                
                # Extrai texto de cada página
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    extracted_content += page.get_text() + "\n"
                    
                    # Processa imagens da página
                    image_list = page.get_images(full=True)
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Processa imagem e adiciona ao Qdrant
                        img_base64 = self.image_handler.encode_image_bytes(image_bytes)
                        description = self.image_handler.image_summarize(img_base64)
                        
                        # Metadata para imagem extraída
                        image_metadata = {
                            "parent_file_id": material_id,
                            "type": "extracted_image",
                            "page_number": page_num + 1,
                            "image_index": img_index + 1
                        }
                        
                        # Adiciona descrição da imagem ao texto extraído
                        extracted_content += f"\nImagem {img_index + 1} da página {page_num + 1}: {description}\n"
                
                pdf_document.close()
                
                # Processa tabelas usando pdfplumber
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        tables = page.extract_tables()
                        for table_index, table in enumerate(tables):
                            if not table:
                                continue
                                
                            df = pd.DataFrame(table[1:], columns=table[0])
                            table_text = df.to_csv(index=False)
                            
                            # Metadata para tabela extraída
                            table_metadata = {
                                "parent_file_id": material_id,
                                "type": "extracted_table",
                                "page_number": page_num + 1,
                                "table_index": table_index + 1
                            }
                            
                            # Adiciona conteúdo da tabela ao texto extraído
                            extracted_content += f"\nTabela {table_index + 1} da página {page_num + 1}:\n{table_text}\n"
                
            elif file_type == 'doc':
                # Processa DOC/DOCX
                doc_file = io.BytesIO(content)
                doc = docx.Document(doc_file)
                
                # Extrai texto dos parágrafos
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        extracted_content += paragraph.text + "\n"
                
                # Extrai texto das tabelas
                for table_index, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    if table_data:
                        df = pd.DataFrame(table_data[1:], columns=table_data[0])
                        table_text = df.to_csv(index=False)
                        extracted_content += f"\nTabela {table_index + 1}:\n{table_text}\n"
                
            elif file_type == 'image':
                # Processa imagem
                image = Image.open(io.BytesIO(content))
                
                # Realiza OCR se a imagem tiver texto
                if image.mode in ('RGBA', 'LA'):
                    background = Image.new('RGB', image.size, (255, 255, 255))
                    background.paste(image, mask=image.split()[-1])
                    image = background
                
                ocr_text = pytesseract.image_to_string(image, lang='por+eng')
                if ocr_text.strip():
                    extracted_content += f"Texto detectado na imagem:\n{ocr_text}\n"
                
                # Gera descrição da imagem
                img_base64 = self.image_handler.encode_image_bytes(content)
                description = self.image_handler.image_summarize(img_base64)
                extracted_content += f"\nDescrição da imagem:\n{description}\n"
            
            # Limpa e normaliza o texto extraído
            cleaned_text = self._clean_text(extracted_content)
            return cleaned_text
            
        except Exception as e:
            print(f"[ERROR] Erro na extração de texto: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    def _clean_text(self, text: str) -> str:
        """Limpa e normaliza o texto extraído."""
        if not text:
            return ""
            
        # Remove caracteres especiais e espaços extras
        cleaned = " ".join(text.split())
        
        # Remove linhas vazias extras
        cleaned = "\n".join(line.strip() for line in cleaned.splitlines() if line.strip())
        
        return cleaned

    async def _process_and_store_in_qdrant(self, material: Material, content: bytes):
        """Processa o conteúdo do arquivo e armazena no Qdrant baseado no tipo."""
        try:
            print(f"[PROCESS] Processando material: {material.id}")
            
            # Extrai texto e conteúdo estruturado
            extracted_text = await self._extract_text_content(content, material.type, material.id)
            
            # Calcula hash do conteúdo extraído
            content_hash = hashlib.sha256(extracted_text.encode('utf-8')).hexdigest()
            
            # Verifica se já existe no Qdrant
            if self.qdrant_handler.document_exists(content_hash, material.student_email, material.discipline_id):
                print(f"[PROCESS] Documento já existe no Qdrant: {material.id}")
                return
            
            # Prepara metadados base
            base_metadata = {
                "file_id": material.id,
                "name": material.name,
                "type": material.type,
                "access_level": material.access_level,
                "session_id": material.session_id,
                "content_hash": content_hash
            }
            
            if self.text_splitter and len(extracted_text) > 1000:
                # Divide texto em chunks menores
                chunks = self.text_splitter.split_text(extracted_text)
                
                # Adiciona cada chunk como documento separado
                for chunk_idx, chunk in enumerate(chunks):
                    chunk_metadata = {
                        **base_metadata,
                        "chunk_index": chunk_idx + 1,
                        "total_chunks": len(chunks)
                    }
                    
                    self.qdrant_handler.add_document(
                        student_email=material.student_email,
                        disciplina=material.discipline_id,
                        content=chunk,
                        metadata_extra=chunk_metadata
                    )
                    print(f"[PROCESS] Adicionado chunk {chunk_idx + 1}/{len(chunks)} ao Qdrant")
            else:
                # Adiciona documento inteiro
                self.qdrant_handler.add_document(
                    student_email=material.student_email,
                    disciplina=material.discipline_id,
                    content=extracted_text,
                    metadata_extra=base_metadata
                )
                print(f"[PROCESS] Material adicionado ao Qdrant: {material.id}")
            
        except Exception as e:
            print(f"[ERROR] Erro ao processar material {material.id}: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
